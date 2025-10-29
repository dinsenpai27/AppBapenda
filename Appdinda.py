import os
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ------------------------
def hapus_border_tabel(table):
    tbl = table._element
    tblPr = tbl.xpath(".//w:tblPr")
    if not tblPr:
        return
    tblPr = tblPr[0]
    tblBorders = tblPr.xpath(".//w:tblBorders")
    if tblBorders:
        tblPr.remove(tblBorders[0])
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def proses(folder_gambar, folder_output, nama_word):
    try:
        os.makedirs(folder_output, exist_ok=True)
        ekstensi_valid = [".jpg", ".jpeg", ".png", ".bmp"]

        if not os.path.exists(folder_gambar):
            messagebox.showerror("Error", f"❌ Folder tidak ditemukan:\n{folder_gambar}")
            return

        gambar_list = [os.path.join(folder_gambar, f)
                       for f in os.listdir(folder_gambar)
                       if os.path.splitext(f)[1].lower() in ekstensi_valid]
        gambar_list.sort()

        if not gambar_list:
            messagebox.showwarning("Peringatan", "Tidak ada gambar yang ditemukan di folder ini.")
            return

        doc = Document()
        section = doc.sections[0]
        section.page_height = Inches(13)
        section.page_width = Inches(8.5)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        for i in range(0, len(gambar_list), 4):
            batch = gambar_list[i:i+4]
            table = doc.add_table(rows=2, cols=2)
            table.autofit = False
            hapus_border_tabel(table)

            for j, gambar_path in enumerate(batch):
                row = j // 2
                col = j % 2
                cell = table.cell(row, col)
                paragraf = cell.paragraphs[0]
                paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraf.add_run()

                try:
                    img = Image.open(gambar_path)
                    if img.width > img.height:
                        img = img.rotate(90, expand=True)

                    tmp_path = os.path.join(folder_output, "tmp_rotated.png")
                    img.save(tmp_path)
                    run.add_picture(tmp_path, width=Inches(2.93), height=Inches(5.69))
                except Exception as e:
                    print(f"Gagal menyisipkan {gambar_path}: {e}")

            if i + 4 < len(gambar_list):
                doc.add_page_break()

        output_file = os.path.join(folder_output, f"{nama_word}.docx")
        doc.save(output_file)
        messagebox.showinfo("Selesai ✅", f"Dokumen berhasil dibuat:\n{output_file}")
    except Exception as e:
        messagebox.showerror("Error", f"Terjadi kesalahan:\n{e}")


# ------------------------
# GUI Aplikasi
root = tk.Tk()
root.title("AppDinda - Konversi Gambar ke Word")
root.geometry("500x300")

tk.Label(root, text="Folder Gambar:").pack(pady=3)
entry_gambar = tk.Entry(root, width=60)
entry_gambar.pack()
tk.Button(root, text="Pilih", command=lambda: entry_gambar.insert(0, filedialog.askdirectory())).pack(pady=2)

tk.Label(root, text="Folder Output:").pack(pady=3)
entry_output = tk.Entry(root, width=60)
entry_output.pack()
tk.Button(root, text="Pilih", command=lambda: entry_output.insert(0, filedialog.askdirectory())).pack(pady=2)

tk.Label(root, text="Nama File Word:").pack(pady=3)
entry_word = tk.Entry(root, width=60)
entry_word.pack()

tk.Button(root, text="Proses", bg="#4CAF50", fg="white",
          command=lambda: proses(entry_gambar.get(), entry_output.get(), entry_word.get())).pack(pady=15)

root.mainloop()
